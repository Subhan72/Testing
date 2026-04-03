"""
HTML and DOC generation for RAG-friendly output.
Produces: transcript.html, explanation.html, transcript.docx, explanation.docx.
Uses inline/basic CSS. Explanation includes heading tags and table of contents.
"""
import os
import re
from typing import Dict, List, Optional
import html as html_lib

import config
import cv2
import numpy as np
from PIL import Image
from api_config import APIConfig
try:
    from google.genai import Client as _GenaiClient
except Exception:
    _GenaiClient = None

# Relative path from output_dir to enhanced images folder (same directory level as HTML files)
ENHANCED_IMAGES_SUBDIR = "enhanced_diagrams"


def format_timestamp(seconds: float) -> str:
    """Format seconds as MM:SS or HH:MM:SS."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


# Basic inline CSS for RAG-friendly structure
BASE_CSS = """
body { font-family: Georgia, serif; font-size: 11pt; line-height: 1.5; color: #222; max-width: 800px; margin: 0 auto; padding: 1em; }
h1 { font-size: 1.6em; margin-top: 0; border-bottom: 1px solid #ccc; padding-bottom: 0.3em; }
h2 { font-size: 1.3em; margin-top: 1.2em; }
h3 { font-size: 1.1em; margin-top: 1em; }
nav.toc { margin: 1em 0; padding: 1em; background: #f8f8f8; border: 1px solid #e0e0e0; }
nav.toc ul { list-style: none; padding-left: 0; }
nav.toc li { margin: 0.4em 0; }
nav.toc a { color: #0066cc; text-decoration: none; }
section.segment { margin: 1em 0; padding-bottom: 0.8em; border-bottom: 1px solid #eee; }
span.timestamp { font-size: 0.9em; color: #666; }
figure { margin: 1em 0; text-align: center; }
figure img { max-width: 50%; height: auto; }
figcaption { font-size: 0.9em; color: #666; margin-top: 0.3em; }
.MathJax { font-size: 1.1em; }
"""

# MathJax 3: render $...$ (inline) and $$...$$ (display) in explanation HTML
MATHJAX_SCRIPT = """
<script>
  window.MathJax = {
    tex: {
      inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
      displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']],
      processEscapes: true
    },
    options: {
      skipHtmlTags: ['script', 'noscript', 'style', 'textarea', 'pre']
    }
  };
</script>
<script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js" async></script>
"""


def _split_math_segments(text: str):
    """
    Split text into segments: (is_math, mode, content).
    Matches $$...$$ (display) and $...$ (inline). Content for math is LaTeX without delimiters.
    """
    pattern = re.compile(r"\$\$[\s\S]*?\$\$|\$[^\$]*?\$")
    segments = []
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            segments.append((False, text[last_end : m.start()]))
        raw = m.group(0)
        if raw.startswith("$$") and raw.endswith("$$"):
            segments.append((True, "display", raw[2:-2].strip()))
        else:
            segments.append((True, "inline", raw[1:-1].strip()))
        last_end = m.end()
    if last_end < len(text):
        segments.append((False, text[last_end:]))
    return segments


def _paragraph_to_html_with_math(para: str) -> str:
    """Convert paragraph with **bold** and $/$$ math to HTML. Math is left as $...$ / $$...$$ for MathJax to render."""
    para = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", para)
    segments = _split_math_segments(para)
    out = []
    for seg in segments:
        if not seg[0]:
            escaped = html_lib.escape(seg[1])
            escaped = escaped.replace("&lt;strong&gt;", "<strong>").replace("&lt;/strong&gt;", "</strong>")
            out.append(escaped)
        else:
            mode, latex = seg[1], seg[2]
            # Escape < and > so they don't break HTML; MathJax will read the text
            latex_safe = latex.replace("<", "&lt;").replace(">", "&gt;")
            if mode == "display":
                out.append(f"$${latex_safe}$$")
            else:
                out.append(f"${latex_safe}$")
    return "<p>" + "".join(out) + "</p>"


def _normalize_doc_text(text: str) -> str:
    """
    Normalize markdown/math-rich text for DOCX plain paragraph rendering.
    DOCX output here is text-first (not true equation objects), so we strip
    delimiters and keep readable math content.
    """
    if not text:
        return ""
    s = text.strip()
    # Markdown bold/italic markers
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"__(.+?)__", r"\1", s)
    # Math delimiters: $...$, $$...$$, \( ... \), \[ ... \]
    s = re.sub(r"\$\$([\s\S]*?)\$\$", r"\1", s)
    s = re.sub(r"\$([^\$]*?)\$", r"\1", s)
    s = re.sub(r"\\\((.*?)\\\)", r"\1", s)
    s = re.sub(r"\\\[(.*?)\\\]", r"\1", s)
    # Common escaped characters from markdown/latex text
    s = s.replace(r"\_", "_").replace(r"\%", "%").replace(r"\&", "&")
    # Basic LaTeX symbol aliases for readability
    replacements = {
        r"\times": "x",
        r"\cdot": "·",
        r"\leq": "<=",
        r"\geq": ">=",
        r"\neq": "!=",
        r"\rightarrow": "->",
        r"\to": "->",
    }
    for k, v in replacements.items():
        s = s.replace(k, v)
    return s


class HtmlDocGenerator:
    """Generates HTML and DOCX for transcript and explanation (RAG-friendly)."""

    # Step 2.5 digital diagrams: fixed max display (non-dynamic). Hand-drawn uses OpenCV sizing.
    DIGITAL_HTML_MAX_PCT = 100
    DIGITAL_DOC_MAX_IN = 6.5  # ≈ full text column on Letter with 1" margins

    def __init__(self):
        self.enhanced_subdir = ENHANCED_IMAGES_SUBDIR
        self._ai_alt_cache: Dict[str, str] = {}
        self._gemini_client = None
        # Keep model aligned with diagram_enhancer.py for image understanding tasks.
        self._gemini_model_name = "gemini-3-pro-image-preview"
        if _GenaiClient is not None:
            try:
                self._gemini_client = _GenaiClient(api_key=APIConfig.get_google_api_key())
            except Exception:
                self._gemini_client = None

    def _estimate_visual_scale(self, image_path: str) -> float:
        """
        Estimate how "large" text/diagram strokes appear in an image.
        Higher score => visually larger content.
        Returns a normalized score in [0, 1].
        """
        try:
            img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
            if img is None:
                return 0.5

            h, w = img.shape[:2]
            if h < 32 or w < 32:
                return 0.5

            # Robust binarization for mixed whiteboard/blackboard outputs.
            blur = cv2.GaussianBlur(img, (5, 5), 0)
            bw = cv2.adaptiveThreshold(
                blur,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY_INV,
                35,
                11,
            )

            # Foreground occupancy correlates with "zoomed-in / oversized" outputs.
            fg_ratio = float(np.count_nonzero(bw)) / float(bw.size)

            # Character/shape size proxy from connected components.
            num_labels, _, stats, _ = cv2.connectedComponentsWithStats(bw, connectivity=8)
            heights = []
            areas = []
            min_area = max(16, int(0.00002 * h * w))
            max_area = int(0.25 * h * w)
            for i in range(1, num_labels):  # skip background
                x, y, ww, hh, area = stats[i]
                if area < min_area or area > max_area:
                    continue
                # Ignore thin artifacts/lines.
                if ww <= 1 or hh <= 1:
                    continue
                heights.append(float(hh))
                areas.append(float(area))

            if not heights:
                return min(max(fg_ratio * 2.0, 0.0), 1.0)

            median_h = float(np.median(np.array(heights)))
            p90_h = float(np.percentile(np.array(heights), 90))
            median_h_norm = median_h / float(h)
            p90_h_norm = p90_h / float(h)

            # Weighted score: larger component heights + higher ink density => larger perceived size.
            score = (
                0.45 * min(max(median_h_norm / 0.045, 0.0), 1.0)
                + 0.30 * min(max(p90_h_norm / 0.09, 0.0), 1.0)
                + 0.25 * min(max(fg_ratio / 0.20, 0.0), 1.0)
            )
            return float(min(max(score, 0.0), 1.0))
        except Exception:
            return 0.5

    def _dynamic_image_widths(self, image_path: str) -> tuple[int, float]:
        """
        Convert visual scale score into target render sizes for HTML and DOCX.
        Returns:
            (html_max_width_percent, doc_width_inches)
        """
        score = self._estimate_visual_scale(image_path)

        # Base widths in your current templates.
        html_base = 50.0
        doc_base = 5.0

        # If score is high (oversized text/diagram), reduce width.
        # If low, allow slightly larger display.
        # Clamp to safe reading bounds.
        shrink = (score - 0.5) * 0.9  # approx [-0.45, +0.45]
        html_width = html_base * (1.0 - shrink)
        doc_width = doc_base * (1.0 - shrink)

        html_width = int(round(min(max(html_width, 45.0), 90.0)))
        # Match DOC bounds to the same 40%-90% scaling intent of the base width (5.0in).
        doc_width = float(min(max(doc_width, 2.5), 5.0))
        return html_width, doc_width

    def _is_step3_dynamic_eligible(self, metadata: Optional[Dict]) -> bool:
        """
        Apply dynamic sizing only for Step 3 enhanced hand-drawn diagrams.
        Step 2.5 digital diagrams must keep fixed size.
        """
        if not metadata:
            return False
        return (metadata.get("content_type") or "").strip().lower() != "digital"

    def _diagram_by_timestamp(
        self,
        diagram_metadata: List[Dict],
        enhanced_diagram_paths: Dict[int, str],
    ) -> Dict[float, List[Dict]]:
        """Map timestamp -> list of {diagram_id, path, metadata} for transcript."""
        out = {}
        for d in diagram_metadata:
            ts = d.get("timestamp", 0)
            did = d.get("diagram_id")
            if did is None or did not in enhanced_diagram_paths:
                continue
            path = enhanced_diagram_paths[did]
            if not path or not os.path.exists(path):
                continue
            if ts not in out:
                out[ts] = []
            out[ts].append({"diagram_id": did, "path": path, "metadata": d})
        return out

    def _relative_image_path(self, absolute_path: str, output_dir: str) -> str:
        """Return path relative to output_dir for use in HTML (e.g. enhanced_diagrams/enhanced_diagram_0001.png or diagrams/diagram_0001.png)."""
        try:
            rel = os.path.relpath(absolute_path, output_dir)
            return rel.replace("\\", "/")
        except ValueError:
            return os.path.basename(absolute_path)

    def _build_image_alt_text(self, metadata: Optional[Dict], fallback_label: str = "diagram") -> str:
        """
        Build accessible, descriptive alt text from diagram metadata.
        """
        if not metadata:
            return fallback_label

        did = metadata.get("diagram_id")
        ts = metadata.get("timestamp")
        board_type = metadata.get("board_type")
        content_type = metadata.get("content_type")

        desc_parts = []
        desc_parts.append("lecture diagram")

        if content_type:
            if str(content_type).lower() == "digital":
                desc_parts.append("digital")
            elif str(content_type).lower() == "hand_drawn":
                desc_parts.append("hand-drawn")

        if board_type and str(board_type).lower() != "unknown":
            desc_parts.append(f"from {board_type}")

        if did is not None:
            desc_parts.append(f"id {did}")

        if ts is not None:
            try:
                desc_parts.append(f"at {format_timestamp(float(ts))}")
            except Exception:
                pass

        # Keep the result concise and readable for screen readers.
        return ", ".join(desc_parts) if desc_parts else fallback_label

    def _is_generic_alt_text(self, alt_text: str) -> bool:
        """Detect low-information generic captions and force a retry."""
        if not alt_text:
            return True
        s = alt_text.strip().lower()
        generic_starts = (
            "a block diagram",
            "this diagram",
            "the diagram shows",
            "a diagram illustrates",
            "an image of",
        )
        if any(s.startswith(p) for p in generic_starts):
            return True
        if len(s.split()) < 6:
            return True
        return False

    def _build_alt_text_with_ai(
        self,
        image_path: str,
        metadata: Optional[Dict],
        fallback_label: str,
    ) -> str:
        """
        Build descriptive alt text using Gemini from image only, with safe fallback.
        """
        # Stable cache key: prefer diagram_id; fallback to absolute path.
        cache_key = str(metadata.get("diagram_id")) if metadata and metadata.get("diagram_id") is not None else image_path
        if cache_key in self._ai_alt_cache:
            return self._ai_alt_cache[cache_key]

        fallback_alt = self._build_image_alt_text(metadata, fallback_label=fallback_label)
        if self._gemini_client is None or not os.path.exists(image_path):
            self._ai_alt_cache[cache_key] = fallback_alt
            return fallback_alt

        try:
            ts = float((metadata or {}).get("timestamp", 0.0) or 0.0)
            content_type = (metadata or {}).get("content_type", "")
            board_type = (metadata or {}).get("board_type", "")
            did = (metadata or {}).get("diagram_id", "unknown")
            prompt = f"""Create a concise, accessibility-focused description of what this diagram explains.

Constraints:
- Output exactly one sentence, 14-28 words.
- Explain the specific concept/process/relationship visible in this diagram.
- Do not invent facts not visible in the image.
- Avoid generic starts like "A block diagram illustrates", "This diagram", or "The diagram shows".
- Mention concrete elements visible in the diagram (e.g., labels, arrows, blocks, equations, axes, components).

Metadata:
- Diagram id: {did}
- Content type: {content_type}
- Board type: {board_type}
- Timestamp: {format_timestamp(ts)}
"""
            img = Image.open(image_path)
            response = self._gemini_client.models.generate_content(
                model=self._gemini_model_name,
                contents=[prompt, img],
                config={"temperature": 0.1, "max_output_tokens": 120},
            )
            alt_text = (getattr(response, "text", None) or "").strip()
            # Same extraction style robustness as diagram_enhancer: inspect parts if text is empty.
            if not alt_text and getattr(response, "parts", None):
                for part in response.parts:
                    if getattr(part, "text", None):
                        alt_text = part.text.strip()
                        if alt_text:
                            break
            alt_text = re.sub(r"\s+", " ", alt_text).strip().strip('"')
            # Retry once with stricter instruction if output is generic.
            if self._is_generic_alt_text(alt_text):
                retry_prompt = (
                    prompt
                    + "\n\nReturn one specific sentence only. Start with the main topic directly, "
                      "for example: 'Flow of ...', 'Signal path for ...', 'Architecture of ...'."
                )
                response = self._gemini_client.models.generate_content(
                    model=self._gemini_model_name,
                    contents=[retry_prompt, img],
                    config={"temperature": 0.05, "max_output_tokens": 120},
                )
                alt_text = (getattr(response, "text", None) or "").strip()
                if not alt_text and getattr(response, "parts", None):
                    for part in response.parts:
                        if getattr(part, "text", None):
                            alt_text = part.text.strip()
                            if alt_text:
                                break
                alt_text = re.sub(r"\s+", " ", alt_text).strip().strip('"')
            if not alt_text:
                alt_text = fallback_alt
            if self._is_generic_alt_text(alt_text):
                # Final fallback to non-generic metadata-derived description.
                alt_text = fallback_alt
            # Hard cap for safety/readability in alt attributes.
            if len(alt_text) > 220:
                alt_text = alt_text[:217].rstrip() + "..."
            self._ai_alt_cache[cache_key] = alt_text
            return alt_text
        except Exception:
            self._ai_alt_cache[cache_key] = fallback_alt
            return fallback_alt

    def generate_transcript_html(
        self,
        output_path: str,
        transcript_data: Dict,
        diagram_metadata: List[Dict],
        enhanced_diagram_paths: Dict[int, str],
        video_name: Optional[str] = None,
        output_dir: Optional[str] = None,
    ) -> str:
        """Generate transcript HTML with segments and images at timestamps. RAG-friendly."""
        output_dir = output_dir or os.path.dirname(output_path)
        diagram_by_ts = self._diagram_by_timestamp(diagram_metadata, enhanced_diagram_paths)
        segments = transcript_data.get("segments", [])
        title = (video_name or "Lecture") + " — Transcript"

        parts = [
            "<!DOCTYPE html>",
            "<html lang=\"en\">",
            "<head>",
            "<meta charset=\"UTF-8\">",
            "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">",
            f"<title>{html_lib.escape(title)}</title>",
            "<style>",
            BASE_CSS,
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{html_lib.escape(title)}</h1>",
            "<p>Full transcript with diagrams at the timestamps where they appear in the lecture.</p>",
        ]

        for seg in segments:
            seg_start = seg.get("start", 0)
            seg_end = seg.get("end", 0)
            seg_text = seg.get("text", "").strip()
            if not seg_text:
                continue

            diagram_to_insert = None
            for ts, diagrams in list(diagram_by_ts.items()):
                if seg_start <= ts <= seg_end and diagrams:
                    diagram_to_insert = diagrams[0]
                    diagram_by_ts.pop(ts, None)
                    break

            ts_str = format_timestamp(seg_start)
            parts.append("<section class=\"segment\">")
            parts.append(f"<span class=\"timestamp\">[{ts_str}]</span> ")
            parts.append(f"<span>{html_lib.escape(seg_text)}</span>")

            if diagram_to_insert:
                rel_path = self._relative_image_path(diagram_to_insert["path"], output_dir)
                cap_ts = format_timestamp(diagram_to_insert["metadata"].get("timestamp", 0))
                alt_text = self._build_alt_text_with_ai(
                    diagram_to_insert["path"],
                    diagram_to_insert.get("metadata"),
                    fallback_label=f"Diagram at {cap_ts}",
                )
                if self._is_step3_dynamic_eligible(diagram_to_insert["metadata"]):
                    html_w, _ = self._dynamic_image_widths(diagram_to_insert["path"])
                else:
                    html_w = self.DIGITAL_HTML_MAX_PCT
                parts.append(
                    f"<figure><img src=\"{html_lib.escape(rel_path)}\" "
                    f"style=\"max-width: {html_w}%; height: auto;\" "
                    f"alt=\"{html_lib.escape(alt_text)}\"/><figcaption>Diagram at {cap_ts}</figcaption></figure>"
                )

            parts.append("</section>")

        # Remaining diagrams
        for ts, diagrams in diagram_by_ts.items():
            for di in diagrams:
                rel_path = self._relative_image_path(di["path"], output_dir)
                cap_ts = format_timestamp(di["metadata"].get("timestamp", 0))
                alt_text = self._build_alt_text_with_ai(
                    di["path"],
                    di.get("metadata"),
                    fallback_label=f"Diagram at {cap_ts}",
                )
                if self._is_step3_dynamic_eligible(di["metadata"]):
                    html_w, _ = self._dynamic_image_widths(di["path"])
                else:
                    html_w = self.DIGITAL_HTML_MAX_PCT
                parts.append("<section class=\"segment\">")
                parts.append(f"<span class=\"timestamp\">[{cap_ts}]</span> ")
                parts.append(
                    f"<figure><img src=\"{html_lib.escape(rel_path)}\" "
                    f"style=\"max-width: {html_w}%; height: auto;\" "
                    f"alt=\"{html_lib.escape(alt_text)}\"/><figcaption>Diagram at {cap_ts}</figcaption></figure>"
                )
                parts.append("</section>")

        parts.append("</body></html>")
        html = "\n".join(parts)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
        if config.VERBOSE:
            print(f"Transcript HTML: {output_path}")
        return output_path

    def _explanation_to_html_with_toc(
        self,
        explanation_text: str,
        diagram_metadata: List[Dict],
        enhanced_diagram_paths: Dict[int, str],
        output_dir: str,
        transcript_data: Optional[Dict] = None,
    ) -> tuple:
        """
        Parse explanation text (## / ### and [DIAGRAM:N]), produce (html_body_string, toc_entries).
        toc_entries = [(level, id, text), ...] for building TOC.
        """
        ordered = sorted(diagram_metadata, key=lambda d: d.get("timestamp", 0))
        diagram_index_to_entry = {}
        for i, d in enumerate(ordered, start=1):
            did = d.get("diagram_id")
            if did and did in enhanced_diagram_paths and os.path.exists(enhanced_diagram_paths[did]):
                diagram_index_to_entry[i] = {
                    "path": self._relative_image_path(enhanced_diagram_paths[did], output_dir),
                    "metadata": d,
                }

        # Replace [DIAGRAM:N] with placeholder so we don't break markdown parsing
        pattern = re.compile(r"\[DIAGRAM:(\d+)\]")
        def diagram_placeholder(m):
            n = m.group(1)
            return f"\n\n<<<DIAGRAM:{n}>>>\n\n"
        text = pattern.sub(diagram_placeholder, explanation_text)

        toc_entries = []
        lines = text.split("\n")
        out = []
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Headings: ## -> h2, ### -> h3, # -> h1
            if re.match(r"^#{1,3}\s+", line):
                level = 0
                if line.startswith("###"):
                    level = 3
                elif line.startswith("##"):
                    level = 2
                elif line.startswith("#"):
                    level = 1
                heading_text = re.sub(r"^#+\s*", "", line).strip()
                heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
                if not heading_text:
                    i += 1
                    continue
                slug = re.sub(r"[^\w\s-]", "", heading_text.lower())
                slug = re.sub(r"[-\s]+", "-", slug).strip("-") or "sec"
                idx = 1
                base_slug = slug
                while any(e[1] == slug for e in toc_entries):
                    idx += 1
                    slug = f"{base_slug}-{idx}"
                toc_entries.append((level, slug, heading_text))
                tag = f"h{min(level, 3)}"
                out.append(f"<{tag} id=\"{html_lib.escape(slug)}\">{html_lib.escape(heading_text)}</{tag}>")
                i += 1
                continue

            # Diagram placeholder
            if "<<<DIAGRAM:" in line and ">>>" in line:
                m = re.search(r"<<<DIAGRAM:(\d+)>>>", line)
                if m:
                    n = int(m.group(1))
                    entry = diagram_index_to_entry.get(n)
                    if entry:
                        path = entry["path"]
                        alt_text = self._build_alt_text_with_ai(
                            os.path.join(output_dir, path),
                            entry.get("metadata"),
                            fallback_label=f"Diagram {n}",
                        )
                        if self._is_step3_dynamic_eligible(entry.get("metadata")):
                            html_w, _ = self._dynamic_image_widths(os.path.join(output_dir, path))
                        else:
                            html_w = self.DIGITAL_HTML_MAX_PCT
                        out.append(
                            f"<figure><img src=\"{html_lib.escape(path)}\" "
                            f"style=\"max-width: {html_w}%; height: auto;\" "
                            f"alt=\"{html_lib.escape(alt_text)}\"/><figcaption>Figure {n}</figcaption></figure>"
                        )
                    else:
                        out.append(f"<!-- Diagram {n} (image not found) -->")
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^---+$", stripped):
                out.append("<hr/>")
                i += 1
                continue

            # Paragraph (with $ / $$ math rendered via MathJax)
            if stripped:
                out.append(_paragraph_to_html_with_math(stripped))
            i += 1

        return "\n".join(out), toc_entries

    def generate_explanation_html(
        self,
        output_path: str,
        explanation_text: str,
        diagram_metadata: List[Dict],
        enhanced_diagram_paths: Dict[int, str],
        video_name: Optional[str] = None,
        output_dir: Optional[str] = None,
        transcript_data: Optional[Dict] = None,
    ) -> str:
        """Generate explanation HTML with TOC and heading tags (RAG-friendly)."""
        output_dir = output_dir or os.path.dirname(output_path)
        title = (video_name or "Lecture") + " — Detailed Explanation"
        body_html, toc_entries = self._explanation_to_html_with_toc(
            explanation_text, diagram_metadata, enhanced_diagram_paths, output_dir, transcript_data=transcript_data
        )

        toc_lines = ["<nav class=\"toc\" aria-label=\"Table of contents\">", "<h2>Table of Contents</h2>", "<ul>"]
        for level, sid, text in toc_entries:
            toc_lines.append(f"<li><a href=\"#{html_lib.escape(sid)}\">{html_lib.escape(text)}</a></li>")
        toc_lines.append("</ul></nav>")
        toc_html = "\n".join(toc_lines)

        parts = [
            "<!DOCTYPE html>",
            "<html lang=\"en\">",
            "<head>",
            "<meta charset=\"UTF-8\">",
            "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">",
            f"<title>{html_lib.escape(title)}</title>",
            "<style>",
            BASE_CSS,
            "</style>",
            MATHJAX_SCRIPT.strip(),
            "</head>",
            "<body>",
            f"<h1>{html_lib.escape(title)}</h1>",
            "<p>Detailed explanation of the lecture with figures. Suitable for RAG and study notes.</p>",
            toc_html,
            body_html,
            """<script>
window.addEventListener('load', function() {
  if (window.MathJax && window.MathJax.typesetPromise) {
    window.MathJax.typesetPromise();
  }
});
</script>""",
            "</body></html>",
        ]
        html = "\n".join(parts)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
        if config.VERBOSE:
            print(f"Explanation HTML: {output_path}")
        return output_path

    def generate_transcript_doc(
        self,
        output_path: str,
        transcript_data: Dict,
        diagram_metadata: List[Dict],
        enhanced_diagram_paths: Dict[int, str],
        video_name: Optional[str] = None,
        output_dir: Optional[str] = None,
    ) -> str:
        """Generate transcript as DOCX with clear structure for RAG (title, section heading, timestamped segments, figures)."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            raise ImportError("python-docx is required for DOC output. Install: pip install python-docx")

        output_dir = output_dir or os.path.dirname(output_path)
        diagram_by_ts = self._diagram_by_timestamp(diagram_metadata, enhanced_diagram_paths)
        segments = transcript_data.get("segments", [])
        doc = Document()
        title = (video_name or "Lecture") + " — Transcript"
        doc.add_heading(title, level=0)
        doc.add_paragraph("Full transcript with diagrams at the timestamps where they appear in the lecture. Suitable for RAG.")
        doc.add_heading("Transcript", level=1)

        for seg in segments:
            seg_start = seg.get("start", 0)
            seg_end = seg.get("end", 0)
            seg_text = seg.get("text", "").strip()
            if not seg_text:
                continue
            diagram_to_insert = None
            for ts, diagrams in list(diagram_by_ts.items()):
                if seg_start <= ts <= seg_end and diagrams:
                    diagram_to_insert = diagrams[0]
                    diagram_by_ts.pop(ts, None)
                    break
            ts_str = format_timestamp(seg_start)
            p = doc.add_paragraph(style="Normal")
            run_ts = p.add_run(f"[{ts_str}] ")
            run_ts.italic = True
            run_ts.font.size = Pt(10)
            p.add_run(seg_text)
            if diagram_to_insert and os.path.exists(diagram_to_insert["path"]):
                try:
                    if self._is_step3_dynamic_eligible(diagram_to_insert["metadata"]):
                        _, doc_w = self._dynamic_image_widths(diagram_to_insert["path"])
                    else:
                        doc_w = self.DIGITAL_DOC_MAX_IN
                    doc.add_picture(diagram_to_insert["path"], width=Inches(doc_w))
                    cap_ts = format_timestamp(diagram_to_insert["metadata"].get("timestamp", 0))
                    doc.add_paragraph(f"Figure at {cap_ts}", style="Caption")
                except Exception:
                    if config.VERBOSE:
                        print("Warning: Could not add diagram image to DOC")

        for ts, diagrams in diagram_by_ts.items():
            for di in diagrams:
                if os.path.exists(di["path"]):
                    try:
                        doc.add_heading("Additional diagram", level=2)
                        if self._is_step3_dynamic_eligible(di["metadata"]):
                            _, doc_w = self._dynamic_image_widths(di["path"])
                        else:
                            doc_w = self.DIGITAL_DOC_MAX_IN
                        doc.add_picture(di["path"], width=Inches(doc_w))
                        doc.add_paragraph(f"Figure at {format_timestamp(di['metadata'].get('timestamp', 0))}", style="Caption")
                    except Exception:
                        pass

        doc.save(output_path)
        if config.VERBOSE:
            print(f"Transcript DOC: {output_path}")
        return output_path

    def generate_explanation_doc(
        self,
        output_path: str,
        explanation_text: str,
        diagram_metadata: List[Dict],
        enhanced_diagram_paths: Dict[int, str],
        video_name: Optional[str] = None,
        output_dir: Optional[str] = None,
    ) -> str:
        """Generate explanation as DOCX with TOC, Heading 1/2/3 hierarchy, and figures (RAG-friendly)."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            raise ImportError("python-docx is required for DOC output. Install: pip install python-docx")

        output_dir = output_dir or os.path.dirname(output_path)
        ordered = sorted(diagram_metadata, key=lambda d: d.get("timestamp", 0))
        diagram_index_to_entry = {}
        for i, d in enumerate(ordered, start=1):
            did = d.get("diagram_id")
            if did and did in enhanced_diagram_paths:
                diagram_index_to_entry[i] = {
                    "path": enhanced_diagram_paths[did],
                    "metadata": d,
                }

        doc = Document()
        title = (video_name or "Lecture") + " — Detailed Explanation"
        doc.add_heading(title, level=0)
        doc.add_paragraph("Detailed explanation with figures. Structured with headings and table of contents for RAG.")

        # Insert a real Word TOC field (updates in Word: right-click TOC -> Update field).
        doc.add_heading("Table of Contents", level=1)
        toc_paragraph = doc.add_paragraph()
        run = toc_paragraph.add_run()
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            fld_char_begin = OxmlElement("w:fldChar")
            fld_char_begin.set(qn("w:fldCharType"), "begin")

            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = r'TOC \o "1-3" \h \z \u'

            fld_char_separate = OxmlElement("w:fldChar")
            fld_char_separate.set(qn("w:fldCharType"), "separate")

            placeholder = OxmlElement("w:t")
            placeholder.text = "Right-click and Update Field to generate Table of Contents."

            fld_char_end = OxmlElement("w:fldChar")
            fld_char_end.set(qn("w:fldCharType"), "end")

            run._r.append(fld_char_begin)
            run._r.append(instr_text)
            run._r.append(fld_char_separate)
            run._r.append(placeholder)
            run._r.append(fld_char_end)
        except Exception:
            # Fallback: keep doc generation robust if field insertion fails.
            toc_paragraph.add_run("Table of Contents will populate from Heading styles.")
        doc.add_paragraph()

        doc.add_heading("Explanation", level=1)

        pattern = re.compile(r"\[DIAGRAM:(\d+)\]")
        parts = pattern.split(explanation_text)
        i = 0
        while i < len(parts):
            if i % 2 == 0 and parts[i].strip():
                block = parts[i]
                for line in block.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if re.match(r"^###\s+", line):
                        heading_text = _normalize_doc_text(re.sub(r"^#+\s*", "", line).strip())
                        doc.add_heading(heading_text, level=3)
                    elif re.match(r"^##\s+", line):
                        heading_text = _normalize_doc_text(re.sub(r"^#+\s*", "", line).strip())
                        doc.add_heading(heading_text, level=2)
                    elif re.match(r"^#\s+", line):
                        heading_text = _normalize_doc_text(re.sub(r"^#+\s*", "", line).strip())
                        doc.add_heading(heading_text, level=1)
                    elif re.match(r"^---+$", line):
                        pass
                    else:
                        doc.add_paragraph(_normalize_doc_text(line), style="Normal")
            elif i % 2 == 1:
                try:
                    n = int(parts[i])
                    entry = diagram_index_to_entry.get(n)
                    path = entry["path"] if entry else None
                    if path and os.path.exists(path):
                        if self._is_step3_dynamic_eligible(entry.get("metadata")):
                            _, doc_w = self._dynamic_image_widths(path)
                        else:
                            doc_w = self.DIGITAL_DOC_MAX_IN
                        doc.add_picture(path, width=Inches(doc_w))
                        doc.add_paragraph(f"Figure {n}", style="Caption")
                except ValueError:
                    pass
            i += 1

        doc.save(output_path)
        if config.VERBOSE:
            print(f"Explanation DOC: {output_path}")
        return output_path
